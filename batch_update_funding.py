#!/usr/bin/env python3
"""
批量更新 company.html 融资表格数据
"""

import re
from docx import Document

# 读取 company.html
with open('company.html', 'r', encoding='utf-8') as f:
    content = f.read()

# 读取 Word 文档
doc = Document('C:/Users/ZhuanZ/Desktop/公司的融资轮次.docx')

# 公司名映射 (Word文档名 -> company.html中的key)
name_mapping = {
    'Figure AI': 'Figure AI',
    '1X Technologies': '1X Technologies',
    'Skild AI': 'Skild AI',
    'Physical Intelligence': 'Physical Intelligence',
    '智元机器人': '智元机器人',
    '宇树科技': '宇树科技',
    '星尘智能': '星尘智能',
    '银河通用': '银河通用',
    '苏度科技': '苏度科技',
    '星海图': '星海图',
    '至简动力': '至简动力',
    '逐际动力': '逐际动力',
    '普渡机器人': '普渡机器人',
    '灵心巧手': '灵心巧手',
    '优必选': '优必选',
    '它石智航': '它石智航',
    '智平方': '智平方',
    '千寻智能': '千寻智能',
    '自变量机器人': '自变量机器人',
    '魔法原子': '魔法原子',
    '乐聚机器人': '乐聚机器人',
    'Sunday Robotics': 'Sunday Robotics',
    '傅利叶智能': '傅利叶',
    'Agility Robotics': 'Agility Robotics',
    'Boston Dynamics': 'Boston Dynamics',
    '思灵机器人': '思灵机器人',
    '小鹏鹏行': '小鹏鹏行',
    '自然意志': '自然意志',
    'Field AI': 'Field AI',
    '梅卡曼德': '梅卡曼德',
    '破壳机器人': '破壳机器人',
    '灵初智能': '灵初智能',
    '珞石机器人': '珞石机器人',
    '地瓜机器人': '地瓜机器人',
    '觅蜂科技': '觅蜂科技',
    '大晓机器人': '大晓机器人',
    '七腾机器人': '七腾机器人',
    '云深处': '云深处',
    '简智机器人': '简智机器人',
    '戴盟机器人': '戴盟机器人',
    '镜识科技': '镜识科技',
    '优理奇智能': '优理奇智能',
    '松延动力': '松延动力',
    '开普勒人形机器人': '开普勒人形机器人',
    '理工华汇': '理工华汇',
    '智在无界': '智在无界',
    '卓益得机器人': '卓益得机器人',
    '天链机器人': '天链机器人',
    '青瞳视觉': '青瞳视觉',
    '钛虎机器人': '钛虎机器人',
    '爱动超越': '爱动超越',
    '灵宇宙': '灵宇宙',
    '灵御智能': '灵御智能',
    '超维动力': '超维动力',
    '无界动力': '无界动力',
    'Anybotics': 'Anybotics',
    'Mimic Robotics': 'Mimic Robotics',
    '加速进化': '加速进化',
    '穹彻智能': '穹彻智能',
}

# 从Word文档提取公司名和表格
def extract_company_tables(doc):
    companies = {}
    current_company = None
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        match = re.match(r'^\d+\.\s*(.+)（', text)
        if match:
            current_company = match.group(1).strip()
            # 找下一个表格
            for j in range(i+1, min(i+10, len(doc.paragraphs))):
                if doc.paragraphs[j].text.strip() == '表格':
                    if j+1 < len(doc.paragraphs):
                        try:
                            table = doc.tables[len([c for c in companies.values() if c is not None]) + 1]
                            # 重新查找：数一下前面有多少表格
                            table_count = 0
                            for t_idx, t in enumerate(doc.tables):
                                if t_idx == table_count:
                                    companies[current_company] = t
                                    break
                                table_count += 1
                        except:
                            pass
                    break
    return companies

# 更可靠的方法：直接按顺序匹配
def get_all_tables_in_order(doc):
    """按文档顺序返回所有表格"""
    return list(doc.tables)

# 公司名 -> 表格索引 的映射（从之前输出得到的）
company_table_indices = {
    'Figure AI': 0,
    '1X Technologies': 1,
    'Skild AI': 2,
    'Physical Intelligence': 3,
    '智元机器人': 4,
    '宇树科技': 5,
    '星尘智能': 6,
    '银河通用': 7,
    '苏度科技': 8,
    '星海图': 9,
    '至简动力': 10,
    '逐际动力': 11,
    '普渡机器人': 12,
    '灵心巧手': 13,
    '优必选': 14,
    '它石智航': 15,
    '智平方': 16,
    '千寻智能': 17,
    '自变量机器人': 18,
    '魔法原子': 19,
    '乐聚机器人': 20,
    'Sunday Robotics': 21,
    '傅利叶': 22,
    'Agility Robotics': 23,
    'Boston Dynamics': 24,
    '思灵机器人': 25,
    '小鹏鹏行': 26,
    '自然意志': 27,
    'Field AI': 28,
    '梅卡曼德': 29,
    '破壳机器人': 30,
    '灵初智能': 31,
    '珞石机器人': 32,
    '地瓜机器人': 33,
    '觅蜂科技': 34,
    '大晓机器人': 35,
    '七腾机器人': 36,
    '云深处': 37,
    '简智机器人': 38,
    '戴盟机器人': 39,
    '镜识科技': 40,
    '优理奇智能': 41,
    '松延动力': 42,
    '开普勒人形机器人': 43,
    '理工华汇': 44,
    '智在无界': 45,
    '卓益得机器人': 46,
    '天链机器人': 47,
    '青瞳视觉': 48,
    '钛虎机器人': 49,
    '爱动超越': 50,
    '灵宇宙': 51,
    '灵御智能': 52,
    '超维动力': 53,
    '无界动力': 54,
    'Anybotics': 55,
    'Mimic Robotics': 56,
    '加速进化': 57,
    '穹彻智能': 58,
}

def escape_js(s):
    if s is None:
        return '—'
    return s.replace('\\', '\\\\').replace("'", "\\'").replace('\n', ' ').replace('\r', '').strip() or '—'

def table_to_funding_data(table):
    """将表格转换为 fundingTable 格式"""
    if not table or len(table.rows) < 2:
        return None
    
    rows = []
    for row in table.rows[1:]:  # 跳过表头
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 5:
            rows.append({
                'round': escape_js(cells[0]),
                'date': escape_js(cells[1]),
                'amount': escape_js(cells[2]),
                'valuation': escape_js(cells[3]),
                'investors': escape_js(cells[4])
            })
    return rows if rows else None

# 获取所有表格
all_tables = list(doc.tables)

# 更新计数器
updated = 0
not_found = []

for company_name, table_idx in company_table_indices.items():
    if table_idx >= len(all_tables):
        not_found.append(f"{company_name}: 表格{table_idx}不存在")
        continue
    
    funding_data = table_to_funding_data(all_tables[table_idx])
    if not funding_data:
        not_found.append(f"{company_name}: 无有效数据")
        continue
    
    # 查找 company.html 中的公司
    # 尝试多种模式
    patterns = [
        rf"('{company_name}':\s*\{{[^}}]*?)(milestones:\s*\[)",
        rf"('{company_name}':\s*\{{[^}}]*?)(investors:\s*\[)",
    ]
    
    replaced = False
    for pattern in patterns:
        match = re.search(pattern, content, re.DOTALL)
        if match:
            # 构建 fundingTable 代码
            table_lines = []
            for entry in funding_data:
                line = f"        {{ round: '{entry['round']}', date: '{entry['date']}', amount: '{entry['amount']}', valuation: '{entry['valuation']}', investors: '{entry['investors']}' }}"
                table_lines.append(line)
            
            new_funding = "fundingTable: [\n" + ",\n".join(table_lines) + "\n        ],"
            
            # 替换 milestones 或 investors
            old_text = match.group(1) + match.group(2)
            new_text = match.group(1) + "fundingTable: [\n" + ",\n".join(table_lines) + "\n        ],\n        "
            
            content = content.replace(old_text, new_text, 1)
            replaced = True
            updated += 1
            print(f"✓ {company_name}")
            break
    
    if not replaced:
        not_found.append(f"{company_name}: 未在company.html中找到")

# 写回文件
with open('company.html', 'w', encoding='utf-8') as f:
    f.write(content)

print(f"\n更新完成: {updated} 家公司")
if not_found:
    print(f"未处理: {len(not_found)} 家")
    for item in not_found[:10]:
        print(f"  - {item}")
